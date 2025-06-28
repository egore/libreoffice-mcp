#!/usr/bin/env python3
"""
Form Letter Generation Module for LibreOffice MCP Server
Handles template processing and mail merge functionality
"""

import os
import subprocess
import tempfile
import uuid
from pathlib import Path
import shutil

# For docx manipulation
from docx import Document


class FormLetterGenerator:
    """
    Class for generating form letters from templates and recipient data
    """

    def __init__(self, libreoffice_path: str = "soffice"):
        """
        Initialize the form letter generator

        Args:
            libreoffice_path: Path to the LibreOffice executable
        """
        self.libreoffice_path = libreoffice_path

    def _ensure_directory(self, directory_path: str) -> None:
        """Ensure the specified directory exists."""
        Path(directory_path).mkdir(parents=True, exist_ok=True)

    def _get_temp_directory(self) -> str:
        """Get a temporary directory for file operations."""
        temp_dir = os.path.join(
            tempfile.gettempdir(), "libreoffice_mcp", str(uuid.uuid4())
        )
        self._ensure_directory(temp_dir)
        return temp_dir

    def _convert_to_pdf(self, input_path: str, output_directory: str = None) -> str:
        """
        Convert a document to PDF using LibreOffice.

        Args:
            input_path: Path to the input document
            output_directory: Directory to save the PDF (optional)

        Returns:
            Path to the output PDF file
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")

        if output_directory is None:
            output_directory = self._get_temp_directory()
        else:
            self._ensure_directory(output_directory)

        # Construct the LibreOffice command
        cmd = [
            self.libreoffice_path,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_directory,
            input_path,
        ]

        try:
            # Run the conversion
            process = subprocess.run(cmd, check=True, capture_output=True, text=True)

            # Get the output filename
            input_filename = os.path.basename(input_path)
            output_filename = os.path.splitext(input_filename)[0] + ".pdf"
            output_path = os.path.join(output_directory, output_filename)

            if not os.path.exists(output_path):
                raise RuntimeError(
                    f"Conversion failed: Output file not found at {output_path}"
                )

            return output_path

        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"LibreOffice conversion failed: {e.stderr}")

    def _replace_placeholders(
        self, template_path: str, recipient_data: dict[str, str]
    ) -> str:
        """
        Replace placeholders in a template with recipient data

        Args:
            template_path: Path to the template document
            recipient_data: Dictionary of field names and values

        Returns:
            Path to the generated document
        """
        # Load the template document
        doc = Document(template_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for field_name, field_value in recipient_data.items():
                placeholder = f"{{{{{field_name}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, field_value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field_name, field_value in recipient_data.items():
                        placeholder = f"{{{{{field_name}}}}}"
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(
                                    placeholder, field_value
                                )

        # Save the generated document
        output_dir = self._get_temp_directory()
        recipient_name = recipient_data.get("name", "recipient").replace(" ", "_")
        output_path = os.path.join(output_dir, f"letter_{recipient_name}.docx")
        doc.save(output_path)

        return output_path

    def generate_form_letters(
        self,
        template_path: str,
        recipients: list[dict[str, str]],
        output_format: str = "pdf",
        output_directory: str = None,
    ) -> list[str]:
        """
        Generate form letters for multiple recipients

        Args:
            template_path: Path to the template document
            recipients: List of dictionaries containing recipient data
            output_format: Output format (pdf or docx)
            output_directory: Directory to save the output files (optional)

        Returns:
            List of paths to the generated documents
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        if output_directory is None:
            output_directory = self._get_temp_directory()
        else:
            self._ensure_directory(output_directory)

        output_paths = []

        for recipient in recipients:
            # Generate the document with replaced placeholders
            docx_path = self._replace_placeholders(template_path, recipient)

            if output_format.lower() == "pdf":
                # Convert to PDF
                output_path = self._convert_to_pdf(docx_path, output_directory)
            else:
                # Copy the DOCX to the output directory
                output_filename = os.path.basename(docx_path)
                output_path = os.path.join(output_directory, output_filename)
                shutil.copy2(docx_path, output_path)

            output_paths.append(output_path)

        return output_paths
