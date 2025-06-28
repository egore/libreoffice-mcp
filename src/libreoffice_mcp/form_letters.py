"""
Form letter generation functionality
"""

import os
import shutil
import logging
from typing import Any
from docx import Document

from .util import get_temp_directory, ensure_directory
from .convert_pdf import PdfConverter

class FormLetterGenerator:
    logger = logging.getLogger(__name__)

    """
    Class for generating form letters from templates and recipient data
    """

    def __init__(self, libreoffice_path: str = "soffice"):
        """
        Initialize the form letter generator

        Args:
            libreoffice_path: Path to the LibreOffice executable
        """
        self.libreoffice_path = libreoffice_path
        self.pdf_converter = PdfConverter(libreoffice_path)

    def _replace_placeholders(
        self, template_path: str, recipient_data: dict[str, str]
    ) -> str:
        """
        Replace placeholders in a template with recipient data

        Args:
            template_path: Path to the template document
            recipient_data: Dictionary of field names and values

        Returns:
            Path to the generated document
        """
        # Load the template document
        doc = Document(template_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for field_name, field_value in recipient_data.items():
                placeholder = f"{{{{{field_name}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, field_value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field_name, field_value in recipient_data.items():
                        placeholder = f"{{{{{field_name}}}}}"
                        for paragraph in cell.paragraphs:
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(
                                    placeholder, field_value
                                )

        # Save the generated document
        output_dir = get_temp_directory()
        recipient_name = recipient_data.get("name", "recipient").replace(" ", "_")
        output_path = os.path.join(output_dir, f"letter_{recipient_name}.docx")
        doc.save(output_path)

        return output_path

    def _generate_form_letters(
        self,
        template_path: str,
        recipients: list[dict[str, str]],
        output_format: str = "pdf",
        output_directory: str = None,
    ) -> list[str]:
        """
        Generate form letters for multiple recipients

        Args:
            template_path: Path to the template document
            recipients: List of dictionaries containing recipient data
            output_format: Output format (pdf or docx)
            output_directory: Directory to save the output files (optional)

        Returns:
            List of paths to the generated documents
        """

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        if output_directory is None:
            output_directory = get_temp_directory()
        else:
            ensure_directory(output_directory)

        output_paths = []

        for recipient in recipients:
            # Generate the document with replaced placeholders
            docx_path = self._replace_placeholders(template_path, recipient)

            if output_format.lower() == "pdf":
                # Convert to PDF
                output_path = self.pdf_converter.convert_to_pdf(docx_path, output_directory)
            else:
                # Copy the DOCX to the output directory
                output_filename = os.path.basename(docx_path)
                output_path = os.path.join(output_directory, output_filename)
                shutil.copy2(docx_path, output_path)

            output_paths.append(output_path)

        return output_paths

    # Define the generate_form_letters tool
    def tool_generate_form_letters(
        self, template_path: str, recipients: list[dict[str, str]], output_format: str = "pdf"
    ) -> dict[str, Any]:
        """Generate form letters from a template and recipient data."""
        try:
            output_paths = self._generate_form_letters(
                template_path=template_path,
                recipients=recipients,
                output_format=output_format,
            )

            return {
                "output_paths": output_paths,
                "success": True,
                "message": f"Successfully generated {len(output_paths)} form letters",
            }
        except Exception as e:
            return {"output_paths": [], "success": False, "message": str(e)}
