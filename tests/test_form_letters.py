#!/usr/bin/env python3
"""
Tests for the form letter generation functionality
"""

import os
import unittest
import tempfile
import shutil

# Add parent directory to path to import modules
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.form_letters import FormLetterGenerator


class TestFormLetterGeneration(unittest.TestCase):
    """Test cases for form letter generation"""

    def setUp(self):
        """Set up test environment"""
        # Create a temporary directory for test files
        self.test_dir = tempfile.mkdtemp()

        # Create a test template file
        self.template_path = os.path.join(self.test_dir, "template.docx")
        self._create_test_template()

        # Create test recipient data
        self.recipients = [
            {
                "name": "John Doe",
                "address": "123 Main St",
                "city": "Anytown",
                "reference": "ABC123",
            },
            {
                "name": "Jane Smith",
                "address": "456 Oak Ave",
                "city": "Somewhere",
                "reference": "XYZ789",
            },
        ]

        # Initialize the form letter generator
        self.generator = FormLetterGenerator()

    def tearDown(self):
        """Clean up test environment"""
        # Remove the temporary directory and its contents
        shutil.rmtree(self.test_dir)

    def _create_test_template(self):
        """Create a test template file using python-docx"""
        try:
            from docx import Document

            # Create a template document with placeholders
            doc = Document()
            doc.add_heading("Form Letter Template", 0)
            doc.add_paragraph("Dear {{name}},")
            doc.add_paragraph("We are writing regarding your account {{reference}}.")
            doc.add_paragraph("Your address on file is:")
            doc.add_paragraph("{{address}}")
            doc.add_paragraph("{{city}}")
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph("The Company")

            # Save the document
            doc.save(self.template_path)
        except ImportError:
            # If python-docx is not available, create an empty file
            # (the test will be skipped)
            with open(self.template_path, "w") as f:
                f.write("Test template")

    def test_generator_initialization(self):
        """Test that the FormLetterGenerator initializes correctly"""
        self.assertIsInstance(self.generator, FormLetterGenerator)
        self.assertEqual(self.generator.libreoffice_path, "soffice")

    def test_replace_placeholders(self):
        """Test placeholder replacement in a template"""
        # Skip if python-docx is not available
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not available")

        # Replace placeholders for the first recipient
        try:
            output_path = self.generator._replace_placeholders(
                self.template_path, self.recipients[0]
            )

            # Check that the output file exists
            self.assertTrue(os.path.exists(output_path))

            # Check that the output file has a .docx extension
            self.assertTrue(output_path.endswith(".docx"))

            # Check that the placeholders were replaced
            doc = Document(output_path)
            content = "\n".join([p.text for p in doc.paragraphs])

            # Check for replaced content
            self.assertIn("Dear John Doe,", content)
            self.assertIn("your account ABC123", content)
            self.assertIn("123 Main St", content)
            self.assertIn("Anytown", content)

            # Check that placeholders were removed
            self.assertNotIn("{{name}}", content)
            self.assertNotIn("{{reference}}", content)
            self.assertNotIn("{{address}}", content)
            self.assertNotIn("{{city}}", content)
        except Exception as e:
            self.fail(f"_replace_placeholders raised an exception: {e}")

    def test_generate_form_letters_docx(self):
        """Test generation of form letters in DOCX format"""
        # Skip if python-docx is not available
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not available")

        # Generate form letters in DOCX format
        try:
            output_paths = self.generator.generate_form_letters(
                template_path=self.template_path,
                recipients=self.recipients,
                output_format="docx",
                output_directory=os.path.join(self.test_dir, "output_docx"),
            )

            # Check that we got the expected number of output files
            self.assertEqual(len(output_paths), len(self.recipients))

            # Check that all output files exist
            for path in output_paths:
                self.assertTrue(os.path.exists(path))
                self.assertTrue(path.endswith(".docx"))
        except Exception as e:
            self.fail(f"generate_form_letters raised an exception: {e}")

    def test_generate_form_letters_pdf(self):
        """Test generation of form letters in PDF format"""
        # Skip if LibreOffice is not available
        if not self._is_libreoffice_available():
            self.skipTest("LibreOffice not available")

        # Generate form letters in PDF format
        try:
            output_paths = self.generator.generate_form_letters(
                template_path=self.template_path,
                recipients=self.recipients,
                output_format="pdf",
                output_directory=os.path.join(self.test_dir, "output_pdf"),
            )

            # Check that we got the expected number of output files
            self.assertEqual(len(output_paths), len(self.recipients))

            # Check that all output files exist
            for path in output_paths:
                self.assertTrue(os.path.exists(path))
                self.assertTrue(path.endswith(".pdf"))
        except Exception as e:
            self.fail(f"generate_form_letters raised an exception: {e}")

    def test_nonexistent_template(self):
        """Test with a nonexistent template file"""
        with self.assertRaises(FileNotFoundError):
            self.generator.generate_form_letters(
                template_path="/nonexistent/template.docx", recipients=self.recipients
            )

    def _is_libreoffice_available(self):
        """Check if LibreOffice is available on the system"""
        import subprocess

        try:
            # Try to run LibreOffice with --version
            subprocess.run(
                ["soffice", "--version"],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            return True
        except (subprocess.SubprocessError, FileNotFoundError):
            return False


if __name__ == "__main__":
    unittest.main()
