#!/usr/bin/env python3
"""
Tests for the DOCX to PDF conversion functionality
"""

import os
import unittest
import tempfile
import shutil

# Add parent directory to path to import modules
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.server import convert_to_pdf


class TestDocxToPdfConversion(unittest.TestCase):
    """Test cases for DOCX to PDF conversion"""

    def setUp(self):
        """Set up test environment"""
        # Create a temporary directory for test files
        self.test_dir = tempfile.mkdtemp()

        # Create a simple test DOCX file
        self.test_docx_path = os.path.join(self.test_dir, "test_document.docx")
        self._create_test_docx()

    def tearDown(self):
        """Clean up test environment"""
        # Remove the temporary directory and its contents
        shutil.rmtree(self.test_dir)

    def _create_test_docx(self):
        """Create a test DOCX file using python-docx"""
        try:
            from docx import Document

            # Create a simple document
            doc = Document()
            doc.add_heading("Test Document", 0)
            doc.add_paragraph("This is a test document for DOCX to PDF conversion.")
            doc.add_paragraph(
                "It contains some simple text to verify the conversion works."
            )

            # Save the document
            doc.save(self.test_docx_path)
        except ImportError:
            # If python-docx is not available, create an empty file
            # (the test will be skipped)
            with open(self.test_docx_path, "w") as f:
                f.write("Test document")

    def test_convert_to_pdf_exists(self):
        """Test that the convert_to_pdf function exists"""
        self.assertTrue(callable(convert_to_pdf))

    def test_convert_to_pdf_output_directory(self):
        """Test conversion with a specified output directory"""
        # Skip if LibreOffice is not available
        if not self._is_libreoffice_available():
            self.skipTest("LibreOffice not available")

        # Create output directory
        output_dir = os.path.join(self.test_dir, "output")
        os.makedirs(output_dir, exist_ok=True)

        # Convert the test document
        try:
            output_path = convert_to_pdf(self.test_docx_path, output_dir)

            # Check that the output file exists
            self.assertTrue(os.path.exists(output_path))

            # Check that the output file is in the specified directory
            self.assertEqual(os.path.dirname(output_path), output_dir)

            # Check that the output file has a .pdf extension
            self.assertTrue(output_path.endswith(".pdf"))
        except Exception as e:
            self.fail(f"convert_to_pdf raised an exception: {e}")

    def test_convert_to_pdf_nonexistent_file(self):
        """Test conversion with a nonexistent input file"""
        # Try to convert a nonexistent file
        with self.assertRaises(FileNotFoundError):
            convert_to_pdf("/nonexistent/file.docx")

    def _is_libreoffice_available(self):
        """Check if LibreOffice is available on the system"""
        import subprocess

        try:
            # Try to run LibreOffice with --version
            subprocess.run(
                ["soffice", "--version"],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            return True
        except (subprocess.SubprocessError, FileNotFoundError):
            return False


if __name__ == "__main__":
    unittest.main()
